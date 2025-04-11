import flet as ft
from docx import Document

class Relat:
    def __init__(self, page):
        self.page = page

        # Inicializa o FilePicker
        self.file_picker = ft.FilePicker(on_result=self.on_file_picker_result)
        self.template_path = ""
        # Adiciona o FilePicker ao overlay da página
        if self.file_picker not in self.page.overlay:
                self.page.overlay.append(self.file_picker)
                self.page.update()  # Atualiza a página para aplicar as alterações
    def salvar_relatorio(self, report_data):
        """
        Solicita o local para salvar o relatório e gera o arquivo.
        """
        self.report_data = report_data
        self.file_picker.save_file(dialog_title="Salvar Relatório", file_name="relatorio_gerado.docx")

    def on_file_picker_result(self, e):
        """
        Callback para lidar com o resultado do FilePicker.
        """
        if e.path:
            try:
                # Gera o relatório no caminho selecionado
                self.gerar_relatorio(e.path)
                self.page.open ( ft.SnackBar(ft.Text(f"Relatório salvo em: {e.path}"), bgcolor=ft.colors.GREEN)) 
            except Exception as ex:
                self.page.open(ft.SnackBar(ft.Text(f"Erro ao salvar relatório: {ex}"), bgcolor=ft.colors.RED)) 
            
            self.page.update()

    def gerar_relatorio(self, output_path):
        """
        Gera o relatório no caminho especificado.
        """
        # Caminho para o modelo .docx
         

        # Carrega o modelo
        doc = Document(self.template_path)

        # Substitui os placeholders pelos valores do relatório
        for paragrafo in doc.paragraphs:
            for chave, valor in self.report_data.items():
                if chave in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(chave, str(valor))

        # Salva o documento no caminho especificado
        doc.save(output_path)

class RelatOcorrencia(Relat):
    def __init__(self, page):
        super().__init__(page)
        self.documento = Document()
        self.template_path = "../docs/ocorrencia.docx"
        self.output_path = "SistemaEscolar/src/model/ModeloOcorrencia.docx"


class RelatMatricula(Relat):
    def __init__(self, page):
        super().__init__(page)
        self.documento = Document()
        self.template_path = "SistemaEscolar/docs/DECLARA_matricula.docx"
        self.output_path = "SistemaEscolar/docs/DECLARAÇÃO_matricula1.docx"

class RelatFrequencia(Relat):
    def __init__(self, page):
        super().__init__(page)
        self.documento = Document()
        self.template_path = "SistemaEscolar/docs/DECLARA_matricula.docx"
        self.output_path = "SistemaEscolar/src/model/ModeloOcorrencia.docx"


class RelatOcorrencia(Relat):
    def __init__(self, page):
        super().__init__(page)
        self.documento = Document()
        self.template_path = "SistemaEscolar/docs/ocorrencia.docx"
        self.output_path = "SistemaEscolar/src/model/ModeloOcorrencia.docx"